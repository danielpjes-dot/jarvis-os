"""
JARVIS Skill — Document Editor
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, Optional


try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None


try:
    from docx import Document
except Exception:
    Document = None


SKILL_NAME = "document_editor"
SKILL_DESCRIPTION = (
    "Read, analyze, summarize, and improve PDF, Word, TXT, Markdown documents "
    "such as CVs, resumes, reports, and letters."
)

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "document_editor",
            "description": SKILL_DESCRIPTION,
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "analyze", "improve_cv", "rewrite", "create_docx"],
                    },
                    "file_path": {"type": "string"},
                    "output_path": {"type": "string"},
                    "content": {"type": "string"},
                    "mode": {"type": "string"},
                    "title": {"type": "string"},
                },
                "required": ["action"],
            },
        },
    }
]


KEYWORDS = {
    "document_editor": [
        "document",
        "pdf",
        "docx",
        "word",
        "cv",
        "resume",
        "read file",
        "analyze document",
        "improve cv",
        "rewrite cv",
        "create docx",
    ]
}


SKILL_META = {
    "route": "reason",
    "keywords": KEYWORDS["document_editor"],
    "tools": {
        "document_editor": {
            "route": "reason",
            "intent_aliases": [
                "read document",
                "analyze pdf",
                "improve cv",
                "rewrite document",
                "create word document",
            ],
            "direct_match": [
                "read this pdf",
                "analyze this document",
                "make my cv better",
                "create docx",
            ],
        }
    },
}


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def _ok(**kwargs: Any) -> Dict[str, Any]:
    return {"ok": True, **kwargs}


def _err(message: str, **kwargs: Any) -> Dict[str, Any]:
    return {"ok": False, "error": message, **kwargs}


def read_pdf(path: Path) -> str:
    if PdfReader is None:
        raise ImportError("PDF support missing. Install with: pip install pypdf")

    reader = PdfReader(str(path))
    parts = []

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"\n--- Page {i} ---\n{text}")

    return "\n".join(parts).strip()


def read_docx(path: Path) -> str:
    if Document is None:
        raise ImportError("DOCX support missing. Install with: pip install python-docx")

    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def read_document(file_path: str) -> Dict[str, Any]:
    path = Path(file_path).expanduser().resolve()

    if not path.exists():
        return _err(f"File not found: {path}")

    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return _err(
            f"Unsupported file type: {ext}",
            supported=sorted(SUPPORTED_EXTENSIONS),
        )

    try:
        if ext == ".pdf":
            text = read_pdf(path)
        elif ext == ".docx":
            text = read_docx(path)
        else:
            text = read_text(path)
    except Exception as e:
        return _err(str(e), path=str(path), extension=ext)

    return _ok(
        path=str(path),
        extension=ext,
        characters=len(text),
        text=text,
    )


def create_docx_file(
    output_path: str,
    content: str,
    title: str = "Improved Document",
) -> Dict[str, Any]:
    if Document is None:
        return _err("DOCX support missing. Install with: pip install python-docx")

    path = Path(output_path).expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    if title:
        doc.add_heading(title, level=1)

    for block in content.splitlines():
        line = block.strip()

        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(str(path))

    return _ok(output_path=str(path))


def build_analysis_prompt(text: str, mode: str = "general") -> str:
    if mode == "cv":
        return f"""
You are an expert CV/resume editor.

Analyze the CV below and give:
1. Overall quality score from 1-10
2. Main weaknesses
3. Missing information
4. Improvements for structure
5. Better wording suggestions
6. ATS/search keyword improvements
7. A rewritten improved version

CV TEXT:
{text}
""".strip()

    return f"""
Analyze this document.

Give:
1. Summary
2. Main issues
3. Structure problems
4. Language improvements
5. Concrete rewrite suggestions

DOCUMENT TEXT:
{text}
""".strip()


def run(
    action: str,
    file_path: Optional[str] = None,
    output_path: Optional[str] = None,
    content: Optional[str] = None,
    mode: str = "general",
    title: str = "Improved Document",
) -> Dict[str, Any]:

    if action in {"read", "analyze", "improve_cv", "rewrite"}:
        if not file_path:
            return _err("file_path is required")

        data = read_document(file_path)

        if not data.get("ok"):
            return data

        text = data["text"]

        if action == "read":
            return data

        if action == "improve_cv":
            mode = "cv"

        prompt = build_analysis_prompt(text, mode=mode)

        return _ok(
            path=data["path"],
            extension=data["extension"],
            characters=data["characters"],
            mode=mode,
            prompt_for_model=prompt,
            note=(
                "Send prompt_for_model to the selected LLM, then optionally call "
                "document_editor with action=create_docx."
            ),
        )

    if action == "create_docx":
        if not output_path:
            return _err("output_path is required")
        if not content:
            return _err("content is required")

        return create_docx_file(
            output_path=output_path,
            content=content,
            title=title,
        )

    return _err(f"Unknown action: {action}")


def document_editor(*args: Any, **kwargs: Any) -> Dict[str, Any]:
    """
    ReAct-safe executor.

    Supports:
      document_editor({"action": "read", "file_path": "..."})
      document_editor(action="read", file_path="...")
    """

    if args and isinstance(args[0], dict):
        payload = dict(args[0])
        payload.update(kwargs)
        return run(**payload)

    return run(*args, **kwargs)


TOOL_MAP = {
    "document_editor": document_editor,
}


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("action", choices=["read", "analyze", "improve_cv", "rewrite", "create_docx"])
    parser.add_argument("--file", dest="file_path")
    parser.add_argument("--output", dest="output_path")
    parser.add_argument("--content")
    parser.add_argument("--mode", default="general")
    parser.add_argument("--title", default="Improved Document")

    cli_args = parser.parse_args()

    result = run(
        action=cli_args.action,
        file_path=cli_args.file_path,
        output_path=cli_args.output_path,
        content=cli_args.content,
        mode=cli_args.mode,
        title=cli_args.title,
    )

    print(json.dumps(result, indent=2, ensure_ascii=False))